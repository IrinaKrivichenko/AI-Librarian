import os
import pandas as pd
from docx import Document


def analyze_docx(file_path):
    doc = Document(file_path)
    num_lines = 0
    num_chars = 0
    num_images = len(doc.inline_shapes)
    num_tables = len(doc.tables)

    for para in doc.paragraphs:
        lines = para.text.split('\n')
        num_lines += len(lines)
        num_chars += sum(len(line) for line in lines)

    return num_lines, num_chars, num_images, num_tables


def process_directory(directory):
    data = []

    for dirpath, dirnames, filenames in os.walk(directory):
        for filename in filenames:
            print(filename)
            if filename.endswith(".docx"):
                file_path = os.path.join(dirpath, filename)
                num_lines, num_chars, num_images, num_tables = analyze_docx(file_path)
                data.append({
                    'Filename': filename,
                    'Number of Lines': num_lines,
                    'Number of Characters': num_chars,
                    'Number of Images': num_images,
                    'Number of Tables': num_tables
                })

    return pd.DataFrame(data)


# Установите путь к вашей директории
directory_path = r"./data/PROGRAMMING"
df = process_directory(directory_path)

# Сохраняем DataFrame в CSV файл
df.to_csv('summary.csv', index=False, encoding='utf-8')

print(df)